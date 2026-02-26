"""Word document parser."""
import io
from typing import Dict, Any

from docx import Document

from app.ingestion.parsers.base import DocumentParser, ParsedDocument


class WordParser(DocumentParser):
    """Word document parser (.docx)."""
    
    def parse(self, file_bytes: bytes) -> ParsedDocument:
        """Parse Word document.
        
        Args:
            file_bytes: Word file content
            
        Returns:
            Parsed document
        """
        doc = Document(io.BytesIO(file_bytes))
        
        text_parts = []
        metadata = {}
        
        # Extract document properties
        core_props = doc.core_properties
        metadata['title'] = core_props.title
        metadata['author'] = core_props.author
        metadata['subject'] = core_props.subject
        metadata['created'] = str(core_props.created) if core_props.created else None
        metadata['modified'] = str(core_props.modified) if core_props.modified else None
        
        # Extract text from paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)
        
        # Extract text from tables
        for table_idx, table in enumerate(doc.tables):
            text_parts.append(f"\n--- Table {table_idx + 1} ---\n")
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_text.append(cell_text)
                if row_text:
                    text_parts.append(' | '.join(row_text))
        
        # Combine text
        full_text = '\n'.join(text_parts)
        
        # Clean text
        full_text = self._clean_text(full_text)
        
        # Extract title
        title = metadata.get('title') or self._extract_title(full_text)
        
        # Count words
        word_count = self._count_words(full_text)
        
        return ParsedDocument(
            text=full_text,
            metadata=metadata,
            title=title,
            author=metadata.get('author'),
            page_count=None,  # Word docs don't have fixed pages
            word_count=word_count
        )
